from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


class ExtractionError(RuntimeError):
    pass


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    text: str
    format: str


def _escape_cell(value: str) -> str:
    return " ".join(str(value or "").split()).replace("|", "\\|")


def _table_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = "| " + " | ".join(_escape_cell(value) for value in normalized[0]) + " |"
    separator = "| " + " | ".join("---" for _ in range(width)) + " |"
    data = ["| " + " | ".join(_escape_cell(value) for value in row) + " |" for row in normalized[1:]]
    return "\n".join([header, separator, *data])


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document

        document = Document(str(path))
    except Exception as exc:
        raise ExtractionError(f"Could not open DOCX {path.name}: {exc}") from exc
    parts: list[str] = []
    for paragraph in document.paragraphs:
        value = (paragraph.text or "").strip()
        if not value:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if style_name.lower().startswith("heading"):
            try:
                level = max(1, min(6, int(style_name.split()[-1])))
            except (ValueError, IndexError):
                level = 1
            parts.append(f"{'#' * level} {value}")
        else:
            parts.append(value)
    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        rendered = _table_markdown(rows)
        if rendered:
            parts.append(rendered)
    return "\n\n".join(parts).strip()


def extract_document(path: Path) -> ExtractedDocument:
    source = Path(path)
    suffix = source.suffix.lower()
    try:
        if suffix == ".md":
            return ExtractedDocument(
                text=source.read_text(encoding="utf-8", errors="replace").strip(),
                format="markdown",
            )
        if suffix == ".docx":
            return ExtractedDocument(text=_extract_docx(source), format="docx")
    except OSError as exc:
        raise ExtractionError(f"Could not read {source.name}: {exc}") from exc
    raise ExtractionError(f"Unsupported document type: {source.suffix}")


__all__ = ["ExtractedDocument", "ExtractionError", "extract_document"]

