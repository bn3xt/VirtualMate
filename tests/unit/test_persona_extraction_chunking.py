from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from virtual_mate.ingestion.chunking import chunk_markdown, estimate_tokens, tokenize
from virtual_mate.ingestion.extraction import extract_document
from virtual_mate.persona import PersonaError, PersonaService


def test_persona_load_and_hot_reload(tmp_path: Path) -> None:
    path = tmp_path / "persona.md"
    path.write_text("# Mateo\n\nVamos por partes.", encoding="utf-8")
    service = PersonaService(path)

    first = service.load()
    path.write_text("# Mateo\n\nPrimero, la evidencia.", encoding="utf-8")
    second = service.reload()

    assert "Vamos por partes" in first.text
    assert "Primero, la evidencia" in second.text
    assert second.estimated_tokens == estimate_tokens(second.text)


@pytest.mark.parametrize("content", ["", "  \n\t"])
def test_persona_rejects_empty_content(tmp_path: Path, content: str) -> None:
    path = tmp_path / "persona.md"
    path.write_text(content, encoding="utf-8")

    with pytest.raises(PersonaError, match="empty"):
        PersonaService(path).load()


def test_markdown_extraction_preserves_text(tmp_path: Path) -> None:
    source = tmp_path / "notes.md"
    source.write_text("# Architecture\n\nRelay Envelope v3.", encoding="utf-8")

    extracted = extract_document(source)

    assert extracted.format == "markdown"
    assert extracted.text == "# Architecture\n\nRelay Envelope v3."


def test_docx_extraction_preserves_headings_paragraphs_and_tables(tmp_path: Path) -> None:
    source = tmp_path / "operations.docx"
    document = Document()
    document.add_heading("Operations", level=1)
    document.add_paragraph("Deploy on Tuesday.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Check"
    table.cell(0, 1).text = "Required"
    table.cell(1, 0).text = "OR-17"
    table.cell(1, 1).text = "Yes"
    document.save(source)

    extracted = extract_document(source)

    assert extracted.format == "docx"
    assert "# Operations" in extracted.text
    assert "Deploy on Tuesday." in extracted.text
    assert "| Check | Required |" in extracted.text
    assert "| OR-17 | Yes |" in extracted.text


def test_chunking_preserves_heading_and_targets_700_with_100_overlap() -> None:
    words = [f"token{i}" for i in range(1_500)]
    markdown = "# Large Section\n\n" + " ".join(words)

    chunks = chunk_markdown(markdown, chunk_size=700, overlap=100)

    assert len(chunks) >= 3
    assert all(chunk.heading == "Large Section" for chunk in chunks)
    assert all(len(tokenize(chunk.text)) <= 700 for chunk in chunks)
    # SentenceSplitter preserves a target overlap while respecting word/sentence
    # boundaries, so one boundary token can legitimately differ.
    assert tokenize(chunks[0].text)[-99:] == tokenize(chunks[1].text)[1:100]
    assert tokenize(chunks[1].text)[-99:] == tokenize(chunks[2].text)[1:100]


def test_chunking_tracks_nearest_nested_heading() -> None:
    chunks = chunk_markdown("# Root\n\nIntro.\n\n## Detail\n\nSpecific fact.")

    assert [chunk.heading for chunk in chunks] == ["Root", "Detail"]
    assert chunks[1].text == "Specific fact."

